"""Извлечение текста ТЗ из документов (шаг 5 конвейера, plan.md §3).

Детерминированный код, БЕЗ LLM: docx / pdf / xlsx / txt → нормализованный текст.
Таблицы рендерятся как pipe-таблицы (| a | b |), чтобы и человек, и LLM видели
структуру строк — в ТЗ характеристики почти всегда лежат таблицей.

Кривые сканы / картинки здесь не обрабатываются — для них нужен OCR
(Tesseract → Claude vision fallback), это отдельный шаг, см. plan.md §3 шаг 5.
"""
from __future__ import annotations

import os
from typing import List


def parse(path: str) -> str:
    """Документ → чистый текст с таблицами. Диспетчер по расширению файла."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        return _parse_docx(path)
    if ext == ".pdf":
        return _parse_pdf(path)
    if ext in (".xlsx", ".xls"):
        return _parse_xlsx(path)
    if ext in (".txt", ".md"):
        with open(path, encoding="utf-8") as f:
            return f.read().strip()
    raise ValueError(
        "Неподдерживаемый формат '%s'. Скан/картинку нужно прогнать через OCR." % ext)


def _render_table(rows: List[List[str]]) -> str:
    """Список строк-ячеек → pipe-таблица. Пустые ячейки сохраняем как ''."""
    out = []
    for row in rows:
        cells = [(" " if c is None else str(c)).replace("\n", " ").strip() for c in row]
        if any(cells):
            out.append("| " + " | ".join(cells) + " |")
    return "\n".join(out)


def _parse_docx(path: str) -> str:
    import docx  # python-docx

    doc = docx.Document(path)
    parts: List[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    for i, table in enumerate(doc.tables):
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        rendered = _render_table(rows)
        if rendered:
            parts.append("[ТАБЛИЦА %d]\n%s" % (i + 1, rendered))
    return "\n\n".join(parts).strip()


def _parse_pdf(path: str) -> str:
    import pdfplumber

    parts: List[str] = []
    with pdfplumber.open(path) as pdf:
        for pageno, page in enumerate(pdf.pages, 1):
            text = page.extract_text() or ""
            if text.strip():
                parts.append(text.strip())
            for i, table in enumerate(page.extract_tables() or []):
                rendered = _render_table(table)
                if rendered:
                    parts.append("[ТАБЛИЦА стр.%d №%d]\n%s" % (pageno, i + 1, rendered))
    text = "\n\n".join(parts).strip()
    if text:
        return text
    # Нет текстового слоя → вероятно скан. Пробуем OCR (шаг 5, ocr.py).
    try:
        from ocr import ocr_pdf, available_backend
        if available_backend() is None:
            raise ValueError(
                "PDF без текстового слоя (скан), а OCR недоступен (нет tesseract/easyocr): %s" % path)
        return ocr_pdf(path)
    except ImportError:
        raise ValueError(
            "PDF без текстового слоя (скан) — нужен OCR (модуль ocr.py): %s" % path)


def _parse_xlsx(path: str) -> str:
    import openpyxl

    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    parts: List[str] = []
    for ws in wb.worksheets:
        rows = [list(r) for r in ws.iter_rows(values_only=True)]
        rendered = _render_table(rows)
        if rendered:
            parts.append("[ЛИСТ '%s']\n%s" % (ws.title, rendered))
    wb.close()
    return "\n\n".join(parts).strip()


if __name__ == "__main__":
    import sys

    print(parse(sys.argv[1]))
