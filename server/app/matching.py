"""Сверка «моё ТЗ ↔ ТЗ закупки» — детерминированный MVP (без LLM).

Извлекает текст из загруженного файла, вытаскивает значимые термины/числовые
требования и считает покрытие. Это первый честный слой; извлечение требований
LLM-ом (Sonnet) подключим позже за тем же интерфейсом — вернёт те же Check/score."""

from __future__ import annotations

import io
import re

from app.schema import Check, MatchResult

_UNIT = r"(?:мм|см|м|кг|г|мл|л|шт|мкм|%|°?c|мпа|вт|квт|в|гц|гост|iso|ту)"
_NUM_REQ = re.compile(rf"[\w().,-]*\s*[≥≤=<>]?\s*\d+[.,]?\d*\s*{_UNIT}\b", re.IGNORECASE)
_WORD = re.compile(r"[а-яёa-z0-9]{4,}", re.IGNORECASE)
_STOP = {
    "поставка", "закупка", "товар", "оказание", "услуги", "выполнение", "работ",
    "нужд", "государственн", "бюджетн", "учреждение", "который", "должен", "также",
    "соответствии", "требования", "наличие", "договор", "контракт",
}


# ---------- извлечение текста из файла ----------

def extract_text(filename: str, content: bytes) -> str:
    name = (filename or "").lower()
    is_binary_format = name.endswith((".docx", ".pdf", ".xlsx", ".xls"))
    try:
        if name.endswith(".docx"):
            return _from_docx(content)
        if name.endswith(".pdf"):
            return _from_pdf(content)
        if name.endswith((".xlsx", ".xls")):
            return _from_xlsx(content)
        return content.decode("utf-8", errors="ignore")
    except Exception:
        # docx/pdf/xlsx — это бинарный ZIP/бинарный формат: если разобрать не
        # получилось (битый файл, пароль и т.п.), decode("utf-8", errors="ignore")
        # даёт не текст, а мусор из обрывков байт — significant_terms() потом
        # находит в этом мусоре случайные «термины» и выдаёт неверный, но
        # уверенный результат сверки вместо честного «не удалось прочитать».
        # Текстовые форматы (.txt и т.п.) так раскодировать можно смело.
        return "" if is_binary_format else content.decode("utf-8", errors="ignore")


def _from_docx(content: bytes) -> str:
    import docx  # python-docx
    doc = docx.Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" ".join(c.text for c in row.cells))
    return "\n".join(parts)


def _from_pdf(content: bytes) -> str:
    import pdfplumber
    out = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            out.append(page.extract_text() or "")
    return "\n".join(out)


def _from_xlsx(content: bytes) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    out = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            out.append(" ".join(str(c) for c in row if c is not None))
    return "\n".join(out)


# ---------- термины ----------

# не ровно «-2 буквы» — иначе «клинки»→«клин» засчитает «клинику» как совпадение.
_RU_ENDINGS = [
    "иями", "иях",
    "ами", "ями", "его", "ому", "ыми", "ими", "ого",
    "ах", "ях", "ов", "ев", "ий", "ый", "их", "ых", "ая", "яя", "ое", "ые", "ие", "ом", "ем", "им", "ым", "ой", "ей",
    "а", "я", "о", "е", "и", "ы", "у", "ю", "й", "ь",
]
_STEM_MIN = 4


def _stem(word: str) -> str:
    w = word.lower()
    if len(w) <= _STEM_MIN:
        return w
    for suf in _RU_ENDINGS:
        if len(w) - len(suf) >= _STEM_MIN and w.endswith(suf):
            return w[: len(w) - len(suf)]
    return w


# «беглая гласная»: клинОк → клинка/клинки теряют «о» перед «к» в косвенных формах
# и множественном числе — суффиксный стеммер этого не видит (буква пропадает ВНУТРИ
# слова). Добавляем вариант без гласной в набор терминов ДОПОЛНИТЕЛЬНО (не взамен) —
# так «клинок» с одной стороны и «клинки» с другой пересекутся по термину «клинк».
_FLEETING_VOWEL_RE = re.compile(r"^(.+[бвгджзйклмнпрстфхцчшщ])[оеё]к$", re.IGNORECASE)


def _fleeting_vowel_variant(word: str) -> str | None:
    m = _FLEETING_VOWEL_RE.match(word)
    return (m.group(1) + "к") if m else None


def significant_terms(text: str) -> set[str]:
    terms: set[str] = set()
    for m in _NUM_REQ.finditer(text):
        terms.add(re.sub(r"\s+", " ", m.group(0).strip().lower()))
    for m in _WORD.finditer(text):
        s = _stem(m.group(0))
        if s not in _STOP and not any(s.startswith(x[:5]) for x in _STOP):
            terms.add(s)
            alt = _fleeting_vowel_variant(s)
            if alt:
                terms.add(alt)
    return terms


# ---------- сравнение ----------

def compare(purchase_text: str, product_text: str, *, product_id: str = "product") -> MatchResult:
    """score = доля значимых терминов ТЗ закупки, покрытых карточкой товара."""
    req = significant_terms(purchase_text)
    have = significant_terms(product_text)
    have_join = " ".join(have)

    if not req:
        return MatchResult(score=0, verdict="eligible_with_gaps",
                           explanation="Не удалось извлечь требования из текста ТЗ закупки "
                                       "(возможно, скан или пустой документ).")

    covered, missing = [], []
    for term in sorted(req):
        if term in have or term in have_join or any(term in h or h in term for h in have):
            covered.append(term)
        else:
            missing.append(term)

    score = round(100 * len(covered) / len(req))
    checks: list[Check] = []
    for t in covered[:8]:
        checks.append(Check(req=t, status="pass"))
    for t in missing[:8]:
        checks.append(Check(req=t, status="gap", note="нет в карточке товара"))

    if score >= 80:
        verdict, expl = "eligible", "Товар покрывает большинство требований ТЗ."
    elif score >= 45:
        verdict, expl = "eligible_with_gaps", "Частичное совпадение — проверьте отмеченные пробелы."
    else:
        verdict, expl = "disqualified", "Совпадение низкое — вероятно, ТЗ про другой товар."

    return MatchResult(score=score, verdict=verdict, checks=checks,
                       explanation=f"{expl} Покрыто {len(covered)} из {len(req)} требований.")
